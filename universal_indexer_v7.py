#!/usr/bin/env python3
"""
Sémantický indexer pro RAG systémy v6 (Ultimate Edition - Zpřesněná taxonomie v2)
- Odstraněno plošné zařazování 'ingest' pod llm_session_history
- Přidána širší a přesnější PATH a FILENAME pravidla (např. Inverter config, Obecne analyzy)
- Systematizovány podkategorie (pro dokumentace a python kód)
- Detekce kódování pomocí chardet
- Extrakce textu z PDF a DOCX
- Deduplikace pomocí SHA-256
- Inteligentní snippety a klíčová slova
"""

import os
import pathlib
import re
import json
import argparse
import hashlib
from datetime import datetime
from collections import Counter
from typing import Tuple, Optional

# --- Externí knihovny pro vylepšenou extrakci ---
try:
    import chardet
    HAS_CHARDET = True
except ImportError:
    print("Upozornění: 'chardet' nenalezen. Detekce kódování bude méně přesná.")
    HAS_CHARDET = False

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    print("Upozornění: 'pdfplumber' nenalezen. PDF se budou indexovat jako binární.")
    HAS_PDFPLUMBER = False

try:
    import docx
    HAS_DOCX = True
except ImportError:
    print("Upozornění: 'python-docx' nenalezen. DOCX se budou indexovat jako binární.")
    HAS_DOCX = False

# ==========================================
# 1. UNIVERZÁLNÍ KONFIGURACE
# ==========================================
TEXT_EXTENSIONS = {'.txt', '.md', '.csv', '.json', '.py', '.html', '.log', '.xml', '.yaml', '.yml'} 
RICH_EXTENSIONS = {'.pdf', '.docx'}

ENCODINGS_FALLBACK = ['utf-8', 'utf-8-sig', 'cp1250', 'latin-1']

EXCLUDE_DIRS = {
    "Archive", "Archive_misc", "Install", 
    "Images", "Foto", "FV_foto", "JK_logs", "solcast_data", 
    "__pycache__", ".git", ".venv", "_RAG_Metadata"
}

EXCLUDE_FILES = {"desktop.ini"}
EXCLUDE_EXTENSIONS = {".sqlite", ".sqlite3", ".db"}

STOP_WORDS_CS_EN = {
    "a", "se", "na", "v", "je", "o", "to", "ze", "s", "z", "k", "do", "i", "pro", "ve",
    "od", "za", "jako", "ale", "byl", "po", "jak", "nebo", "uz", "by", "co", "asi", 
    "the", "and", "in", "to", "of", "for", "is", "on", "with", "as", "by", "at", "an",
    "this", "that", "from", "it", "or", "be", "are", "not", "have", "but", "all", "we"
}

# ==========================================
# 2. NOVÁ TAXONOMIE A KLASIFIKACE (v6)
# ==========================================

PATH_RULES = [
    # IOT hardware
    (r'(?i)(^|[/\\])IOT[/\\]BMS[/\\]', 'bms_telemetry'),
    (r'(?i)(^|[/\\])IOT[/\\]Stridac_', 'inverter_config'),
    (r'(?i)(^|[/\\])IOT[/\\]', 'iot_hardware'),
    
    # GCP cloud
    (r'(?i)(^|[/\\])GCP[/\\]Meteo_scraper_', 'meteo_data_pipeline'),
    (r'(?i)(^|[/\\])GCP[/\\]gcp_miner_project[/\\]', 'data_mining_pipeline'),
    (r'(?i)(^|[/\\])GCP[/\\]Cloud_service_pipeline[/\\]', 'serverless_orchestration'),
    (r'(?i)(^|[/\\])GCP[/\\]', 'cloud_infrastructure'),
    
    # Outpost project
    (r'(?i)(^|[/\\])Outpost[/\\]FV[/\\]', 'solar_energy_system'),
    (r'(?i)(^|[/\\])Outpost[/\\]Geodata[/\\]', 'geospatial_data'),
    (r'(?i)(^|[/\\])Outpost[/\\]Obecne_analyzy[/\\]', 'sociotechnical_analysis'),
    (r'(?i)(^|[/\\])Outpost[/\\]', 'solar_energy_system'),
    
    # Job applications
    (r'(?i)(^|[/\\])Jobs_aplications[/\\]', 'job_application'),
    
    # Github structure
    (r'(?i)(^|[/\\])Github[/\\]Mirror[/\\]', 'github_mirror_data'),
    (r'(?i)(^|[/\\])Github[/\\]', 'github_repository'),
    
    # Skripty local structure
    (r'(?i)(^|[/\\])Skripty_local[/\\]Skripty_analyza_json[/\\]', 'json_analysis_tool'),
    (r'(?i)(^|[/\\])Skripty_local[/\\]Skripty_Eshopy[/\\]', 'ecommerce_scraper'),
    (r'(?i)(^|[/\\])Skripty_local[/\\]Skripty_metodika_RAG[/\\]', 'rag_indexing_tool'),
    (r'(?i)(^|[/\\])Skripty_local[/\\]Skripty_prace[/\\]', 'job_market_scraper'),
    
    # LLM data and Sessions
    (r'(?i)(^|[/\\])Source_raw[/\\]LLM[/\\]', 'llm_session_history'),
    (r'(?i)(^|[/\\])Source_raw[/\\]Json[/\\]', 'llm_session_history'),
    (r'(?i)(^|[/\\])Source_raw[/\\]Outpost_kontext_master[/\\]', 'system_context'),
    
    # Stitch dashboards
    (r'(?i)(^|[/\\])Stitch[/\\]', 'web_dashboard'),
]

FILENAME_RULES = [
    # Session and handoff documents
    (r'(?i)session.*\.json$', 'llm_session_history'),
    (r'(?i)(handoff|injekt).*\.(md|json)$', 'llm_session_history'),
    
    # Analysis & evaluations (catching unstructured JSON files)
    (r'(?i)(semantick[aá]_extrakce|semantick[aá]_analyza|llm_model_evaluation|deterministic_llm_prompting|integrace_diy|iot_dev_ai_coding|kaskadova_detekce_iot|optimalizace|json_general).*\.(json|txt|md)$', 'llm_knowledge_base'),
    
    # User Histories
    (r'(?i)historie(-|_)sledovani.*\.json$', 'user_history'),
    
    # Github Mirror specific (if not caught by path)
    (r'(?i)(commits|tree|profile|repos|index).*\.json$', 'github_mirror_data'),

    # Specifické ingest skripty/dokumenty mimo session
    (r'(?i)gcp_stack_ingest.*\.md$', 'cloud_infrastructure'),
    (r'(?i)fv_ingest.*\.json$', 'configuration_file'),
    
    # Forecast and prediction
    (r'(?i)forecast.*\.csv$', 'timeseries_forecast'),
    (r'(?i)soc_predict.*\.py$', 'source_code_python'),
    
    # BMS and battery
    (r'(?i)bms.*\.(csv|pdf|md)$', 'bms_telemetry'),
    (r'(?i)battery.*\.(csv|md)$', 'bms_telemetry'),
    
    # Scrapers
    (r'(?i)(scraper|miner|sklizec|tezeni).*\.py$', 'web_scraper'),
    
    # Post mortem
    (r'(?i)(post_mortem|pitevni|debug|error_gemini).*\.(md|txt|json)$', 'post_mortem_analysis'),
    
    # CV and resume
    (r'(?i)(cv|resume|životopis).*\.(docx|pdf|md)$', 'cv_resume'),
    
    # Configuration
    (r'(?i)(config|settings|requirements|plugin).*\.(txt|csv|json)$', 'configuration_file'),
    
    # Documentation
    (r'(?i)readme.*\.(md|txt)$', 'documentation_master'),
]

FILENAME_RULES = [
    # Session and handoff documents
    (r'(?i)session.*\.json$', 'llm_session_history'),
    (r'(?i)(handoff|injekt).*\.(md|json)$', 'llm_session_history'),
    
    # Analysis & evaluations (catching unstructured JSON files)
    (r'(?i)(semantick[aá]_extrakce|semantick[aá]_analyza|llm_model_evaluation|deterministic_llm_prompting|integrace_diy|iot_dev_ai_coding|kaskadova_detekce_iot|optimalizace).*\.json$', 'llm_knowledge_base'),
    
    # User Histories
    (r'(?i)historie(-|_)sledovani.*\.json$', 'user_history'),
    
    # Github Mirror specific (if not caught by path)
    (r'(?i)(commits|tree|profile|repos|index).*\.json$', 'github_mirror_data'),

    # Specifické ingest skripty/dokumenty mimo session
    (r'(?i)gcp_stack_ingest.*\.md$', 'cloud_infrastructure'),
    (r'(?i)fv_ingest.*\.json$', 'inverter_config'),
    
    # Forecast and prediction
    (r'(?i)forecast.*\.csv$', 'timeseries_forecast'),
    (r'(?i)soc_predict.*\.py$', 'source_code_python'),
    
    # BMS and battery
    (r'(?i)bms.*\.(csv|pdf|md)$', 'bms_telemetry'),
    (r'(?i)battery.*\.(csv|md)$', 'bms_telemetry'),
    
    # Scrapers
    (r'(?i)(scraper|miner|sklizec|tezeni).*\.py$', 'web_scraper'),
    
    # Post mortem
    (r'(?i)(post_mortem|pitevni|debug|error_gemini).*\.(md|txt|json)$', 'post_mortem_analysis'),
    
    # CV and resume
    (r'(?i)(cv|resume|životopis).*\.(docx|pdf|md)$', 'cv_resume'),
    
    # Configuration
    (r'(?i)(config|settings|requirements|plugin).*\.(txt|csv|json)$', 'configuration_file'),
    
    # Documentation
    (r'(?i)readme.*\.(md|txt)$', 'documentation_master'),
]

EXTENSION_RULES = {
    '.py': 'source_code_python',
    '.json': 'configuration_file',
    '.csv': 'tabular_data',
    '.md': 'documentation',
    '.html': 'web_dashboard',
    '.pdf': 'hardware_datasheet',
    '.docx': 'documentation',
    '.png': 'image_media',
    '.jpg': 'image_media',
    '.webp': 'image_media',
    '.laz': 'lidar_pointcloud',
    '.las': 'lidar_pointcloud',
    '.txt': 'plain_text',
    '.ps1': 'powershell_script',
    '.bat': 'batch_script',
    '.sh': 'shell_script',
}

# Rozšířené subkategorie pro dokumentaci a kód
SUBCATEGORY_RULES = {
    # timeseries & data
    'forecast': 'timeseries_forecast',
    'predikce': 'timeseries_forecast',
    'bms': 'bms_telemetry',
    'meteo': 'weather_data',
    
    # configs & data mapping
    'dockerfile': 'docker_config',
    'requirements': 'python_dependencies',
    'topics.csv': 'mining_topics',
    'categories.csv': 'mining_categories',
    
    # documentation subcategories
    'methodology': 'methodology',
    'metodika': 'methodology',
    'handbook': 'handbook',
    'tutorial': 'user_manual',
    'manual': 'user_manual',
    'technical_specification': 'technical_spec',
    'datasheet': 'technical_spec',
    'analysis': 'analysis_report',
    'analyza': 'analysis_report',
    
    # source_code_python subcategories
    'scraper': 'scraper',
    'sklizec': 'scraper',
    'miner': 'scraper',
    'pipeline': 'pipeline',
    'etl': 'pipeline',
    'util': 'utility',
    'dashboard': 'dashboard',
}

def classify_document(file_path: str) -> Tuple[str, Optional[str], str]:
    """
    Vrátí (hlavní_kategorie, podkategorie, pouzite_pravidlo)
    """
    file_path_lower = file_path.lower()
    file_name = file_path.split('/')[-1].split('\\')[-1]
    file_name_lower = file_name.lower()
    
    # 1. PRAVIDLA PODLE NÁZVU SOUBORU (Nejvyšší priorita)
    for pattern, category in FILENAME_RULES:
        if re.search(pattern, file_name_lower):
            subcategory = None
            for keyword, subcat in SUBCATEGORY_RULES.items():
                if keyword in file_name_lower:
                    subcategory = subcat
                    break
            return category, subcategory, f"FILENAME_RULE: {pattern}"

    # 2. CESTNÍ PRAVIDLA
    for pattern, category in PATH_RULES:
        if re.search(pattern, file_path_lower):
            subcategory = None
            for keyword, subcat in SUBCATEGORY_RULES.items():
                if keyword in file_name_lower:
                    subcategory = subcat
                    break
            return category, subcategory, f"PATH_RULE: {pattern}"
    
    # 3. PRAVIDLA PODLE PŘÍPONY
    ext = '.' + file_name.split('.')[-1].lower() if '.' in file_name else ''
    if ext in EXTENSION_RULES:
        main_category = EXTENSION_RULES[ext]
        
        # 4. ZKONTROLUJ PODKATEGORII
        subcategory = None
        for keyword, subcat in SUBCATEGORY_RULES.items():
            if keyword in file_name_lower:
                subcategory = subcat
                break
        
        # Speciální pravidlo pro JSON sessiony
        if ext == '.json' and ('session' in file_name_lower or 'handoff' in file_name_lower):
            return 'llm_session_history', None, "SPECIAL_RULE: json_session"
        
        return main_category, subcategory, f"EXTENSION_RULE: {ext}"
    
    # 5. FALLBACK
    return 'neklasifikovano', None, "FALLBACK"


# ==========================================
# 3. STAVEBNÍ BLOKY A POMOCNÉ FUNKCE
# ==========================================

def get_file_hash(file_path: pathlib.Path) -> str:
    sha256 = hashlib.sha256()
    try:
        with open(file_path, "rb") as f:
            while chunk := f.read(8192):
                sha256.update(chunk)
        return sha256.hexdigest()
    except Exception:
        return "error_hash"

def extract_keywords(text: str, max_words: int = 7) -> list:
    if not text or len(text) < 100:
        return []
        
    words = re.findall(r'\b[a-zA-ZáčďéěíňóřšťúůýžÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ]{4,}\b', text.lower())
    filtered_words = [w for w in words if w not in STOP_WORDS_CS_EN]
    
    if not filtered_words:
        return []
        
    word_counts = Counter(filtered_words)
    return [word for word, count in word_counts.most_common(max_words)]

def generate_smart_snippet(text: str, max_chars: int = 400) -> str:
    if not text:
        return ""
        
    clean_text = re.sub(r'[#*`_~\[\]]', ' ', text)
    clean_text = re.sub(r'\s+', ' ', clean_text).strip()
    
    sentences = re.split(r'(?<=[.!?])\s+', clean_text)
    
    snippet = ""
    for sentence in sentences:
        if len(snippet) + len(sentence) <= max_chars:
            snippet += sentence + " "
        else:
            if not snippet:
                snippet = sentence[:max_chars] + "..."
            break
            
    return snippet.strip()

def read_text_safely(file_path: pathlib.Path, bytes_limit: int | None = 65536) -> tuple[str | None, str | None]:
    try:
        with open(file_path, 'rb') as f:
            raw_data = f.read(bytes_limit) if bytes_limit else f.read()
            
        if not raw_data:
            return "", "utf-8"

        if HAS_CHARDET:
            import chardet
            result = chardet.detect(raw_data)
            encoding = result['encoding']
            confidence = result['confidence']
            
            if encoding and confidence > 0.7:
                try:
                    return raw_data.decode(encoding), encoding
                except UnicodeDecodeError:
                    pass
        
        for enc in ENCODINGS_FALLBACK:
            try:
                content = raw_data.decode(enc)
                return content, enc
            except UnicodeDecodeError:
                continue
                
    except Exception as e:
        return None, f"error: {str(e)}"
        
    return None, "binary_unreadble"

def extract_pdf_text(file_path: pathlib.Path, max_pages: int = 2) -> str:
    if not HAS_PDFPLUMBER:
        return ""
        
    text_content = ""
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            pages_to_read = min(max_pages, len(pdf.pages))
            for i in range(pages_to_read):
                page_text = pdf.pages[i].extract_text()
                if page_text:
                    text_content += page_text + "\n"
        return text_content
    except Exception:
        return ""

def extract_docx_text(file_path: pathlib.Path, max_paragraphs: int = 10) -> str:
    if not HAS_DOCX:
        return ""
        
    text_content = ""
    try:
        import docx
        doc = docx.Document(str(file_path))
        paragraphs_to_read = min(max_paragraphs, len(doc.paragraphs))
        for i in range(paragraphs_to_read):
            text_content += doc.paragraphs[i].text + "\n"
        return text_content
    except Exception:
        return ""

# ==========================================
# 4. HLAVNÍ RAG MIGRATION PIPELINE V6
# ==========================================

def run_ingest_index(target_path: str):
    source_dir = pathlib.Path(target_path).resolve()
    if not source_dir.exists() or not source_dir.is_dir():
        print(f"Chyba: Cesta {source_dir} neexistuje nebo není složka.")
        return

    rag_context_dir = source_dir / "_RAG_Metadata"
    rag_context_dir.mkdir(exist_ok=True)
    
    metadata_json_file = rag_context_dir / "08_rag_metadata_v7_smart.json"

    print(f"-> Zahajuji RAG indexaci (v7) v: {source_dir}")
    print(f"-> Precizní PATH a FILENAME klasifikace, Smart Snippets, Chardet")
    
    rag_metadata = {
        "index_timestamp": datetime.now().isoformat(),
        "source_directory": str(source_dir),
        "total_files_scanned": 0,
        "duplicates_found": 0,
        "documents": []
    }

    seen_hashes = {} 

    for root, dirs, files in os.walk(source_dir):
        exclude_lower = {d.lower() for d in EXCLUDE_DIRS}
        dirs[:] = [d for d in dirs if d.lower() not in exclude_lower and not d.startswith('_')]
        
        for name in files:
            file_path = pathlib.Path(root) / name
            ext = file_path.suffix.lower()
            
            if name.lower() in EXCLUDE_FILES or ext in EXCLUDE_EXTENSIONS:
                continue

            rel_path = str(file_path.relative_to(source_dir)).replace("\\", "/")
            rag_metadata["total_files_scanned"] += 1
            
            file_hash = get_file_hash(file_path)
            is_duplicate = False
            duplicate_of = None
            
            if file_hash in seen_hashes and file_hash != "error_hash":
                is_duplicate = True
                duplicate_of = seen_hashes[file_hash]
                rag_metadata["duplicates_found"] += 1
            else:
                seen_hashes[file_hash] = rel_path

            content = ""
            encoding = "binary"
            snippet = ""
            keywords = []

            if not is_duplicate:
                if ext in TEXT_EXTENSIONS:
                    content, encoding = read_text_safely(file_path)
                    if content is None:
                        content = ""
                        snippet = f"[{ext.upper()} soubor - nečitelné kódování]"
                elif ext == '.pdf' and HAS_PDFPLUMBER:
                    content = extract_pdf_text(file_path)
                    encoding = "pdf_extracted"
                elif ext == '.docx' and HAS_DOCX:
                    content = extract_docx_text(file_path)
                    encoding = "docx_extracted"
                else:
                    snippet = f"[{ext.upper()} soubor - binární obsah nelze číst napřímo]"

                if content:
                    snippet = generate_smart_snippet(content, max_chars=400)
                    if len(content) > 100 and ext not in ['.json', '.csv', '.yaml', '.yml']:
                         keywords = extract_keywords(content)

            # Klasifikace dokumentu s novou logikou v6
            doc_type, subcategory, matched_rule = classify_document(rel_path)

            doc_entry: dict = {
                "file_path": rel_path,
                "document_type": doc_type,
            }
            
            if subcategory:
                doc_entry["subcategory"] = subcategory
            
            doc_entry["matched_rule"] = matched_rule
            doc_entry["size_bytes"] = file_path.stat().st_size
            doc_entry["last_modified"] = datetime.fromtimestamp(file_path.stat().st_mtime).isoformat()
            
            if is_duplicate:
                doc_entry["is_duplicate"] = True
                doc_entry["duplicate_of"] = duplicate_of
                doc_entry["content_snippet"] = "[DUPLICITA - viz originál]"
            else:
                doc_entry["encoding"] = encoding
                doc_entry["content_snippet"] = snippet
                if keywords:
                    doc_entry["keywords"] = keywords

            rag_metadata["documents"].append(doc_entry)

    metadata_json_file.write_text(json.dumps(rag_metadata, ensure_ascii=False, indent=2), encoding="utf-8")
    
    print("-" * 50)
    print("HOTOVO!")
    print(f"Zkontrolováno souborů: {rag_metadata['total_files_scanned']}")
    print(f"Z toho nalezena a indexována duplicitní data: {rag_metadata['duplicates_found']}x")
    print(f"Sémantická RAG databáze uložena: {metadata_json_file}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Univerzální RAG indexer v7 (Dvouúrovňová taxonomie s regex optimalizací).")
    parser.add_argument("path", help="Absolutní cesta k cílové složce")
    args = parser.parse_args()
    
    run_ingest_index(args.path)